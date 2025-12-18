"""
Document Processor Tool

Multi-format document processing with chunking for RAG pipelines.
Supports PDF, DOCX, TXT, MD, HTML, CSV, JSON formats.
"""

import hashlib
import json
import mimetypes
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

# Type hints for optional dependencies
try:
    from pypdf import PdfReader

    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from bs4 import BeautifulSoup

    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False


@dataclass
class DocumentChunk:
    """Represents a chunk of processed document."""

    content: str
    metadata: dict = field(default_factory=dict)
    chunk_index: int = 0
    chunk_id: str = ""

    def __post_init__(self):
        if not self.chunk_id:
            content_hash = hashlib.md5(self.content.encode()).hexdigest()[:8]
            self.chunk_id = (
                f"{self.metadata.get('source', 'unknown')}_{self.chunk_index}_{content_hash}"
            )


@dataclass
class ProcessedDocument:
    """Represents a fully processed document."""

    source: str
    content: str
    chunks: list[DocumentChunk]
    metadata: dict
    processing_time_ms: float


class DocumentProcessor:
    """
    Multi-format document processor with configurable chunking.

    Supports:
    - PDF (requires pypdf)
    - DOCX (requires python-docx)
    - TXT, MD (native)
    - HTML (requires beautifulsoup4)
    - CSV, JSON (native)
    """

    SUPPORTED_FORMATS = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".html": "text/html",
        ".htm": "text/html",
        ".csv": "text/csv",
        ".json": "application/json",
    }

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: Optional[list[str]] = None,
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Target size for each chunk in characters
            chunk_overlap: Number of characters to overlap between chunks
            separators: List of separators for recursive splitting
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", ". ", " ", ""]

    def process_file(self, file_path: str | Path) -> ProcessedDocument:
        """
        Process a document file and return chunks.

        Args:
            file_path: Path to the document file

        Returns:
            ProcessedDocument with content and chunks

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        start_time = datetime.now()
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported format: {suffix}. Supported: {list(self.SUPPORTED_FORMATS.keys())}"
            )

        # Extract content based on format
        content = self._extract_content(path, suffix)

        # Build metadata
        metadata = self._build_metadata(path)

        # Chunk the content
        chunks = self._chunk_text(content, metadata)

        processing_time = (datetime.now() - start_time).total_seconds() * 1000

        return ProcessedDocument(
            source=str(path),
            content=content,
            chunks=chunks,
            metadata=metadata,
            processing_time_ms=processing_time,
        )

    def process_text(
        self,
        text: str,
        source: str = "direct_input",
        metadata: Optional[dict] = None,
    ) -> ProcessedDocument:
        """
        Process raw text and return chunks.

        Args:
            text: Raw text content
            source: Source identifier
            metadata: Additional metadata

        Returns:
            ProcessedDocument with content and chunks
        """
        start_time = datetime.now()

        base_metadata = {
            "source": source,
            "type": "text",
            "ingestion_timestamp": datetime.now().isoformat(),
            "char_count": len(text),
        }
        if metadata:
            base_metadata.update(metadata)

        chunks = self._chunk_text(text, base_metadata)

        processing_time = (datetime.now() - start_time).total_seconds() * 1000

        return ProcessedDocument(
            source=source,
            content=text,
            chunks=chunks,
            metadata=base_metadata,
            processing_time_ms=processing_time,
        )

    def _extract_content(self, path: Path, suffix: str) -> str:
        """Extract text content from file based on format."""
        if suffix == ".pdf":
            return self._extract_pdf(path)
        elif suffix == ".docx":
            return self._extract_docx(path)
        elif suffix in (".txt", ".md"):
            return path.read_text(encoding="utf-8")
        elif suffix in (".html", ".htm"):
            return self._extract_html(path)
        elif suffix == ".csv":
            return self._extract_csv(path)
        elif suffix == ".json":
            return self._extract_json(path)
        else:
            raise ValueError(f"No extractor for format: {suffix}")

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF file."""
        if not HAS_PYPDF:
            raise ImportError(
                "pypdf is required for PDF processing. Install with: pip install pypdf"
            )

        reader = PdfReader(path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)

    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX file."""
        if not HAS_DOCX:
            raise ImportError(
                "python-docx is required for DOCX processing. Install with: pip install python-docx"
            )

        doc = DocxDocument(path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(paragraphs)

    def _extract_html(self, path: Path) -> str:
        """Extract text from HTML file."""
        if not HAS_BS4:
            raise ImportError(
                "beautifulsoup4 is required for HTML processing. Install with: pip install beautifulsoup4"
            )

        html_content = path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()

        return soup.get_text(separator="\n", strip=True)

    def _extract_csv(self, path: Path) -> str:
        """Extract text from CSV file."""
        import csv

        rows = []
        with open(path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))
        return "\n".join(rows)

    def _extract_json(self, path: Path) -> str:
        """Extract text from JSON file."""
        content = json.loads(path.read_text(encoding="utf-8"))
        return json.dumps(content, indent=2, ensure_ascii=False)

    def _build_metadata(self, path: Path) -> dict:
        """Build metadata dictionary for a file."""
        stat = path.stat()
        return {
            "source": path.name,
            "source_path": str(path.absolute()),
            "file_type": path.suffix.lower(),
            "mime_type": self.SUPPORTED_FORMATS.get(
                path.suffix.lower(), "application/octet-stream"
            ),
            "file_size_bytes": stat.st_size,
            "modified_timestamp": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "ingestion_timestamp": datetime.now().isoformat(),
        }

    def _chunk_text(self, text: str, metadata: dict) -> list[DocumentChunk]:
        """
        Split text into overlapping chunks using recursive character splitting.

        Args:
            text: Text to chunk
            metadata: Base metadata for chunks

        Returns:
            List of DocumentChunk objects
        """
        chunks = []

        # Clean up text
        text = re.sub(r"\n{3,}", "\n\n", text)  # Normalize multiple newlines
        text = text.strip()

        if len(text) <= self.chunk_size:
            chunks.append(
                DocumentChunk(
                    content=text,
                    metadata=metadata.copy(),
                    chunk_index=0,
                )
            )
            return chunks

        # Recursive splitting
        split_texts = self._recursive_split(text, self.separators)

        # Merge small chunks and create overlapping windows
        current_chunk = ""
        chunk_index = 0

        for split_text in split_texts:
            if len(current_chunk) + len(split_text) <= self.chunk_size:
                current_chunk += split_text
            else:
                if current_chunk:
                    chunk_metadata = metadata.copy()
                    chunk_metadata["chunk_index"] = chunk_index
                    chunk_metadata["chunk_char_count"] = len(current_chunk)

                    chunks.append(
                        DocumentChunk(
                            content=current_chunk.strip(),
                            metadata=chunk_metadata,
                            chunk_index=chunk_index,
                        )
                    )
                    chunk_index += 1

                    # Keep overlap from end of current chunk
                    overlap_text = (
                        current_chunk[-self.chunk_overlap :] if self.chunk_overlap > 0 else ""
                    )
                    current_chunk = overlap_text + split_text
                else:
                    current_chunk = split_text

        # Don't forget the last chunk
        if current_chunk.strip():
            chunk_metadata = metadata.copy()
            chunk_metadata["chunk_index"] = chunk_index
            chunk_metadata["chunk_char_count"] = len(current_chunk)

            chunks.append(
                DocumentChunk(
                    content=current_chunk.strip(),
                    metadata=chunk_metadata,
                    chunk_index=chunk_index,
                )
            )

        return chunks

    def _recursive_split(self, text: str, separators: list[str]) -> list[str]:
        """Recursively split text using provided separators."""
        if not separators:
            return [text]

        separator = separators[0]
        remaining_separators = separators[1:]

        if separator:
            splits = text.split(separator)
            # Re-add separator to maintain text integrity
            splits = [s + separator if i < len(splits) - 1 else s for i, s in enumerate(splits)]
        else:
            # Character-level split as last resort
            splits = list(text)

        result = []
        for split in splits:
            if len(split) > self.chunk_size and remaining_separators:
                # Recursively split with next separator
                result.extend(self._recursive_split(split, remaining_separators))
            else:
                result.append(split)

        return result


class DocumentProcessorTool:
    """
    MCP Tool wrapper for DocumentProcessor.

    Provides a standardized interface for document processing operations.
    """

    name = "document_processor"
    description = "Process and chunk documents for RAG pipelines"

    def __init__(self, config: Optional[dict] = None):
        config = config or {}
        self.processor = DocumentProcessor(
            chunk_size=config.get("chunk_size", 1000),
            chunk_overlap=config.get("chunk_overlap", 200),
        )
        self.supported_formats = config.get(
            "supported_formats", ["pdf", "docx", "txt", "md", "html", "csv", "json"]
        )

    async def process(self, file_path: str) -> dict[str, Any]:
        """
        Process a document file.

        Args:
            file_path: Path to the document

        Returns:
            Dictionary with processed document data
        """
        try:
            result = self.processor.process_file(file_path)
            return {
                "success": True,
                "source": result.source,
                "chunk_count": len(result.chunks),
                "total_chars": len(result.content),
                "processing_time_ms": result.processing_time_ms,
                "chunks": [
                    {
                        "id": chunk.chunk_id,
                        "content": chunk.content,
                        "metadata": chunk.metadata,
                    }
                    for chunk in result.chunks
                ],
                "metadata": result.metadata,
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "error_type": type(e).__name__,
            }

    async def process_text(self, text: str, source: str = "direct_input") -> dict[str, Any]:
        """
        Process raw text.

        Args:
            text: Raw text content
            source: Source identifier

        Returns:
            Dictionary with processed document data
        """
        try:
            result = self.processor.process_text(text, source)
            return {
                "success": True,
                "source": result.source,
                "chunk_count": len(result.chunks),
                "total_chars": len(result.content),
                "processing_time_ms": result.processing_time_ms,
                "chunks": [
                    {
                        "id": chunk.chunk_id,
                        "content": chunk.content,
                        "metadata": chunk.metadata,
                    }
                    for chunk in result.chunks
                ],
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "error_type": type(e).__name__,
            }


if __name__ == "__main__":
    # Example usage
    processor = DocumentProcessor(chunk_size=500, chunk_overlap=100)

    sample_text = """
    # Document Processing Guide
    
    This is a sample document that demonstrates the document processing capabilities.
    
    ## Features
    
    The document processor supports multiple formats including PDF, DOCX, TXT, MD, HTML, CSV, and JSON.
    It provides intelligent chunking with configurable overlap for optimal RAG performance.
    
    ## Usage
    
    Simply instantiate the processor and call process_file() or process_text().
    The processor will automatically detect the format and apply appropriate extraction.
    
    ## Chunking Strategy
    
    Uses recursive character splitting with multiple separators:
    1. Paragraph boundaries (double newlines)
    2. Line boundaries (single newlines)  
    3. Sentence boundaries (periods with spaces)
    4. Word boundaries (spaces)
    5. Character-level as last resort
    """

    result = processor.process_text(sample_text, source="example.md")
    print(f"Processed {len(result.chunks)} chunks in {result.processing_time_ms:.2f}ms")
    for chunk in result.chunks:
        print(f"\n--- Chunk {chunk.chunk_index} ({len(chunk.content)} chars) ---")
        print(chunk.content[:200] + "..." if len(chunk.content) > 200 else chunk.content)
